from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs"
OUT.mkdir(exist_ok=True)
DOCX_PATH = OUT / "Informe_Tecnico_Proyecto_Emergencias_Adonis_Alegria.docx"
ARCH_PATH = OUT / "arquitectura_sistema_emergencias.png"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "18794E"
AMBER = "9A6700"
RED = "B42318"
WHITE = "FFFFFF"
BLACK = "111827"


def font_path(bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def create_architecture_diagram():
    image = Image.new("RGB", (1800, 1060), "white")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.truetype(font_path(True), 46) if font_path(True) else ImageFont.load_default()
    box_title = ImageFont.truetype(font_path(True), 30) if font_path(True) else ImageFont.load_default()
    body_font = ImageFont.truetype(font_path(False), 24) if font_path(False) else ImageFont.load_default()
    small_font = ImageFont.truetype(font_path(False), 21) if font_path(False) else ImageFont.load_default()

    draw.text((900, 52), "Arquitectura distribuida del sistema de emergencias", font=title_font, fill="#17365D", anchor="ma")

    def rounded_box(x1, y1, x2, y2, title, lines, fill, outline="#B8C4D1"):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=outline, width=4)
        draw.text(((x1 + x2) / 2, y1 + 34), title, font=box_title, fill="#17365D", anchor="ma")
        y = y1 + 95
        for line in lines:
            draw.text((x1 + 34, y), line, font=body_font, fill="#283747")
            y += 39

    def arrow(x1, y1, x2, y2, label=""):
        draw.line((x1, y1, x2, y2), fill="#2E74B5", width=7)
        angle_x = 18 if x2 >= x1 else -18
        angle_y = 18 if y2 >= y1 else -18
        if abs(x2 - x1) > abs(y2 - y1):
            draw.polygon([(x2, y2), (x2 - angle_x, y2 - 13), (x2 - angle_x, y2 + 13)], fill="#2E74B5")
        else:
            draw.polygon([(x2, y2), (x2 - 13, y2 - angle_y), (x2 + 13, y2 - angle_y)], fill="#2E74B5")
        if label:
            draw.rounded_rectangle(((x1+x2)/2-100, (y1+y2)/2-22, (x1+x2)/2+100, (y1+y2)/2+22), 10, fill="white")
            draw.text(((x1+x2)/2, (y1+y2)/2), label, font=small_font, fill="#2E74B5", anchor="mm")

    rounded_box(90, 175, 520, 470, "Clientes React", ["Ciudadano", "Operador / administrador", "Panel KPI y predicciones", "Leaflet + WebRTC"], "#EAF2F8")
    rounded_box(685, 145, 1130, 510, "API Node.js", ["Express REST", "JWT y Google OAuth", "Servicios y controladores", "Errores y logs", "Socket.IO"], "#EDF7ED")
    rounded_box(1280, 175, 1710, 470, "Servicios externos", ["Google Identity Services", "OpenStreetMap", "OSRM routing", "STUN para WebRTC"], "#FFF7E6")
    rounded_box(160, 690, 610, 965, "Procesamiento", ["Worker Thread", "Clasificacion por CSV", "Notebook de analitica", "Baseline predictivo"], "#F4ECF7")
    rounded_box(745, 690, 1195, 965, "Almacenamiento", ["Map() en memoria", "Usuarios y reportes", "Historial por sala", "JSON analitico/predictivo"], "#FDEDEC")
    rounded_box(1330, 690, 1690, 965, "Tiempo real", ["Salas por emergencia", "Chat y multimedia", "Senalizacion WebRTC", "Eventos de estado"], "#E8F8F5")

    arrow(520, 320, 685, 320, "HTTPS / REST")
    arrow(1130, 320, 1280, 320, "HTTPS")
    arrow(900, 510, 900, 690, "lectura / escritura")
    arrow(760, 510, 500, 690, "tarea paralela")
    arrow(1030, 510, 1450, 690, "Socket.IO")
    arrow(1450, 690, 1150, 500, "eventos")

    draw.text((900, 1025), "Puertos de desarrollo: frontend 5173 | backend 3001", font=small_font, fill="#667085", anchor="mm")
    image.save(ARCH_PATH, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=11, bold=None, color=BLACK, italic=None, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    caption = doc.styles["Caption"]
    caption.font.name = "Arial"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Sistema distribuido de gestión de emergencias"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(hp.runs[0], size=8.5, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("ESPE | Aplicaciones Distribuidas | Página ")
    set_run_font(r, size=8.5, color=MID_GRAY)
    add_field(fp, "PAGE", "1")

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    set_run_font(r)
    return p


def create_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    level.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_number(doc, text, num_id):
    p = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p._p.get_or_add_pPr().append(num_pr)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], header_fill)
        for run in hdr[idx].paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=NAVY)
        hdr[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(run, size=9.2)
            if idx == 0 and len(headers) > 2:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code(doc, code, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(caption)
        set_run_font(r, size=9, bold=True, color=NAVY)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(dedent(code).strip())
    set_run_font(r, size=8.2, color="233044", name="Courier New")
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure_caption(doc, number, title, source=None):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Figura {number}. {title}")
    set_run_font(r, size=9, italic=True, color=MID_GRAY)
    if source:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(8)
        r2 = p2.add_run(f"Nota. {source}")
        set_run_font(r2, size=8.5, italic=True, color=MID_GRAY)


def add_evidence_placeholder(doc, number, title, instruction):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=260, start=240, bottom=260, end=240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESPACIO PARA EVIDENCIA REAL")
    set_run_font(r, size=11, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(instruction)
    set_run_font(r2, size=9, italic=True, color=MID_GRAY)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    set_table_geometry(table, [9360])
    add_figure_caption(doc, number, title, "La captura debe incorporarse después de ejecutar y verificar la funcionalidad descrita.")


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSIDAD DE LAS FUERZAS ARMADAS ESPE")
    set_run_font(r, size=15, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Departamento de Ciencias de la Computación\nCarrera de Ingeniería en Tecnologías de la Información")
    set_run_font(r, size=11, color=MID_GRAY)
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(36)
    p3.paragraph_format.space_after = Pt(14)
    r = p3.add_run("SISTEMA DISTRIBUIDO DE REPORTE,\nGESTIÓN Y ANÁLISIS DE EMERGENCIAS")
    set_run_font(r, size=23, bold=True, color=NAVY)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p4.add_run("Integración de servicios con identidad y trazabilidad")
    set_run_font(r, size=14, italic=True, color=BLUE)
    doc.add_paragraph()
    meta = [
        ("Asignatura", "Aplicaciones Distribuidas"),
        ("Estudiante", "Adonis Alegría"),
        ("Docente", "[Completar nombre del docente]"),
        ("NRC", "[Completar NRC]"),
        ("Periodo académico", "2026"),
        ("Fecha de entrega", "22 de junio de 2026"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        rl = p.add_run(f"{label}: ")
        set_run_font(rl, size=11, bold=True, color=NAVY)
        rv = p.add_run(value)
        set_run_font(rv, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("Sangolquí, Ecuador")
    set_run_font(r, size=10, color=MID_GRAY)
    doc.add_page_break()


def build_document():
    create_architecture_diagram()
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "Índice general", 1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Actualice este índice en Word con Ctrl+A y F9.")
    doc.add_page_break()

    add_heading(doc, "Índice de figuras", 1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Figura"', "Actualice este índice en Word con Ctrl+A y F9.")
    doc.add_page_break()

    add_heading(doc, "Resumen", 1)
    add_body(doc, "El presente informe documenta el desarrollo de un sistema distribuido para el reporte, gestión, seguimiento y análisis de emergencias. La propuesta responde a la necesidad de centralizar la información proporcionada por la ciudadanía, ubicar geográficamente los incidentes y mantener una comunicación inmediata con operadores de policía, bomberos o servicios médicos. La solución utiliza React y Vite en el cliente; Node.js, Express y Socket.IO en el servidor; JWT y Google Identity Services para identidad; Worker Threads para clasificación concurrente de gravedad; Leaflet, OpenStreetMap y OSRM para visualización y trazado de rutas; y notebooks de ciencia de datos para indicadores y predicciones operativas. En el estado actual se encuentran funcionales el registro, el inicio de sesión local, la integración de Google sujeta a la configuración del origen autorizado, la gestión de reportes, las transiciones de estado, el chat, la videollamada grupal, la animación de unidades, los KPI y el módulo predictivo. El almacenamiento de negocio permanece en memoria y constituye la principal limitación de persistencia. Como resultado esperado, se busca consolidar una arquitectura observable, segura y escalable, con persistencia durable, pruebas automatizadas y eventual trazabilidad inmutable de los reportes.")
    p = doc.add_paragraph()
    r1 = p.add_run("Palabras clave: ")
    set_run_font(r1, bold=True)
    r2 = p.add_run("aplicaciones distribuidas, WebSocket, JWT, OAuth 2.0, concurrencia, analítica predictiva.")
    set_run_font(r2)

    add_heading(doc, "1. Introducción", 1)
    add_body(doc, "La atención de emergencias exige que la información circule con rapidez, conserve su integridad y pueda ser consultada simultáneamente por ciudadanos y operadores. Un reporte incompleto, una actualización tardía o una asignación incoherente pueden retrasar la respuesta operativa. Por esta razón, el proyecto plantea una aplicación distribuida que separa la interfaz web, los servicios de negocio, la comunicación en tiempo real, el procesamiento concurrente y el análisis de datos.")
    add_body(doc, "El sistema adopta una arquitectura cliente-servidor. El frontend se ejecuta en el navegador y consume una API REST, mientras que el backend administra identidad, reportes, asignación de unidades, analítica y predicciones. Socket.IO mantiene canales bidireccionales para estados, chat y señalización de videollamadas. De forma complementaria, un Worker Thread clasifica la gravedad de cada reporte sin bloquear el ciclo principal de eventos de Node.js.")
    add_body(doc, "El informe se organiza desde el problema y la justificación hasta la arquitectura, la aplicación de contenidos de la asignatura, el estado comprobable, los resultados y las mejoras pendientes. Las afirmaciones técnicas se relacionan con componentes reales del repositorio y se evita presentar como terminado aquello que aún requiere persistencia, despliegue o evidencia de prueba.")

    add_heading(doc, "2. Planteamiento del problema", 1)
    add_body(doc, "Los canales convencionales de atención dependen principalmente de llamadas telefónicas y de la transferencia verbal de información. Este mecanismo puede saturarse, dificulta adjuntar evidencia multimedia y no siempre permite al ciudadano conocer el estado de su solicitud. También obliga al operador a integrar manualmente ubicación, gravedad, disponibilidad de unidades y comunicación con la persona afectada.")
    add_body(doc, "La problemática afecta a ciudadanos que requieren ayuda, operadores que priorizan recursos y unidades que deben desplazarse al sitio. La ausencia de un flujo digital unificado puede producir registros duplicados, pérdida de contexto, tiempos de respuesta poco visibles y decisiones basadas en información dispersa.")
    add_body(doc, "Una aplicación distribuida aporta separación de responsabilidades, comunicación en tiempo real, autenticación centralizada y capacidad de delegar procesos. Sin embargo, también introduce retos: consistencia entre clientes, control de acceso, tolerancia a fallos, protección de credenciales y persistencia confiable.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Pregunta de desarrollo: ¿cómo integrar identidad, comunicación en tiempo real, procesamiento concurrente, georreferenciación y analítica en una aplicación distribuida que permita gestionar reportes de emergencia de forma coherente, segura y observable?")
    set_run_font(r, bold=True, color=NAVY)

    add_heading(doc, "3. Justificación", 1)
    add_body(doc, "La propuesta es importante porque reúne en un único flujo el registro del incidente, su clasificación, la asignación de una unidad, el seguimiento del estado y la comunicación con el ciudadano. El beneficio principal consiste en reducir la fragmentación de la información y ofrecer una vista operacional compartida.")
    add_body(doc, "Desde el punto de vista tecnológico, el proyecto demuestra la integración de mecanismos distribuidos estudiados en la asignatura: HTTP, Socket.IO, tareas asíncronas, Worker Threads, JWT, Google OAuth 2.0, manejo centralizado de errores, logs, servicios desacoplados y consumo de fuentes externas de mapas y rutas. La solución es viable para fines académicos porque utiliza herramientas abiertas y puede ejecutarse localmente con dos procesos principales.")
    add_body(doc, "El proyecto también crea una base para futuras mejoras. La capa de almacenamiento puede sustituirse por una base de datos transaccional; los eventos pueden publicarse en una cola; y los reportes críticos podrían registrar huellas criptográficas en una red blockchain. Estas extensiones no se consideran implementadas en la versión actual.")

    add_heading(doc, "4. Objetivos del proyecto", 1)
    add_heading(doc, "4.1 Objetivo general", 2)
    add_body(doc, "Desarrollar y documentar una aplicación distribuida para el reporte y gestión de emergencias que integre identidad segura, comunicación en tiempo real, procesamiento concurrente, georreferenciación, observabilidad y análisis de datos.")
    add_heading(doc, "4.2 Objetivos específicos", 2)
    objectives_num = create_decimal_numbering(doc)
    add_number(doc, "Implementar el registro, la autenticación local con JWT y el inicio de sesión de usuarios mediante Google Identity Services, aplicando autorización por roles.", objectives_num)
    add_number(doc, "Integrar reportes, chat, videollamada y cambios de estado mediante API REST y Socket.IO, manteniendo sincronizados a ciudadanos y operadores.", objectives_num)
    add_number(doc, "Procesar la clasificación de gravedad en un Worker Thread y exponer dashboards de indicadores y predicciones que apoyen la planificación operativa.", objectives_num)

    add_heading(doc, "5. Marco teórico", 1)
    add_heading(doc, "5.1 Aplicaciones distribuidas y arquitectura cliente-servidor", 2)
    add_body(doc, "Una aplicación distribuida coordina componentes que se ejecutan en procesos o equipos diferentes y que se comunican mediante una red. En este proyecto, el navegador, el servidor Node.js, Google Identity Services, OSRM y los clientes Socket.IO constituyen participantes independientes. HTTP mantiene el intercambio de recursos y operaciones, mientras que los eventos en tiempo real complementan el modelo solicitud-respuesta (IETF, 2022).")
    add_heading(doc, "5.2 Comunicación WebSocket y Socket.IO", 2)
    add_body(doc, "WebSocket permite establecer un canal persistente y bidireccional entre cliente y servidor. Socket.IO incorpora eventos, reconexión y salas lógicas. En la aplicación, cada emergencia utiliza una sala identificada como emergencia-{id}; allí se distribuyen mensajes, indicadores de escritura y señalización de WebRTC. Esta separación evita enviar información de una emergencia a participantes no relacionados (Socket.IO, 2024).")
    add_heading(doc, "5.3 Concurrencia y Worker Threads", 2)
    add_body(doc, "Node.js resuelve muchas operaciones de entrada y salida mediante un ciclo de eventos. Las tareas intensivas de CPU pueden afectar la capacidad de respuesta si se ejecutan en el hilo principal. El módulo worker_threads permite crear hilos JavaScript paralelos y transferir mensajes entre ellos. En el proyecto, la clasificación de gravedad se ejecuta en un worker independiente que carga reglas desde archivos CSV y devuelve el resultado al servicio principal (Node.js Contributors, 2026).")
    add_heading(doc, "5.4 JWT, OAuth 2.0 e identidad", 2)
    add_body(doc, "JWT se utiliza para transportar afirmaciones firmadas sobre la identidad y el rol del usuario. OAuth 2.0 y OpenID Connect permiten delegar la autenticación a un proveedor confiable. La práctica moderna exige validar estrictamente emisor, audiencia, vigencia y origen autorizado, además de evitar flujos obsoletos o redirecciones ambiguas (IETF, 2025). Google Identity Services entrega una credencial que el backend verifica con el Client ID configurado antes de crear o recuperar una cuenta de tipo usuario (Google, 2025).")
    add_heading(doc, "5.5 Observabilidad, logs y manejo de excepciones", 2)
    add_body(doc, "La observabilidad busca inferir el estado interno del sistema mediante señales como logs, métricas y trazas. El backend implementa un logger Singleton con timestamp, nivel, módulo y datos estructurados. Las excepciones de dominio heredan de AppError y son transformadas por un middleware global en respuestas HTTP consistentes. Esta base puede evolucionar hacia telemetría interoperable con OpenTelemetry (OpenTelemetry Authors, 2025).")
    add_heading(doc, "5.6 Analítica y predicción operativa", 2)
    add_body(doc, "Los KPI sintetizan el comportamiento histórico y actual; no constituyen por sí solos un modelo predictivo. El módulo de predicciones agrega un baseline estadístico interpretable que combina frecuencia histórica por zona, día y hora con una tendencia simple. Su propósito es apoyar la planificación, no automatizar decisiones críticas. La interpretación debe considerar sesgos de cobertura, calidad de datos y diferencia entre detenciones registradas y emergencias futuras.")

    add_heading(doc, "6. Desarrollo del proyecto", 1)
    add_heading(doc, "6.1 Descripción general y usuarios", 2)
    add_body(doc, "El sistema permite a un ciudadano crear reportes con tipo, descripción, ubicación y evidencia. Un operador consulta los incidentes, acepta la petición, indica que la unidad está en camino y marca el caso como resuelto. Los roles implementados son usuario, operador y administrador; los dos últimos acceden a las rutas de gestión, analítica y predicciones.")
    add_table(doc, ["Rol", "Responsabilidad principal", "Acceso"], [
        ("Usuario", "Registrar o consultar reportes y comunicarse en la sala de la emergencia.", "Panel ciudadano, chat y videollamada"),
        ("Operador", "Gestionar reportes, cambiar estados, revisar mapas y analizar indicadores.", "Centro de operaciones, detalle, KPI y predicciones"),
        ("Administrador", "Acceder a las funciones operativas y futuras tareas de administración.", "Rutas habilitadas para operador/administrador"),
    ], [1500, 4300, 3560])

    add_heading(doc, "6.2 Tecnologías utilizadas", 2)
    add_table(doc, ["Capa", "Tecnología", "Aplicación en el proyecto"], [
        ("Frontend", "React, Vite, React Router", "SPA, navegación, contextos de identidad y vistas por rol."),
        ("Backend", "Node.js y Express", "API REST, servicios, controladores y middlewares."),
        ("Tiempo real", "Socket.IO y WebRTC", "Eventos, chat, salas y videollamada grupal."),
        ("Seguridad", "JWT, bcryptjs, Google Identity", "Sesión firmada, hash de contraseñas e identidad federada."),
        ("Mapas", "Leaflet, OpenStreetMap, OSRM", "Ubicación, rutas reales y animación de la unidad."),
        ("Datos", "Pandas/Jupyter, JSON y CSV", "KPI, baseline predictivo y reglas de gravedad."),
        ("Almacenamiento", "Map() en memoria", "Usuarios, reportes e historial temporal; no es persistencia durable."),
    ], [1450, 2650, 5260])

    add_heading(doc, "6.3 Estructura lógica del repositorio", 2)
    add_bullet(doc, "backend/src/controladores: adaptación de solicitudes HTTP y respuestas.")
    add_bullet(doc, "backend/src/servicios: autenticación, reportes, asignación, analítica y predicciones.")
    add_bullet(doc, "backend/src/almacenamiento: repositorios en memoria y unidades de emergencia.")
    add_bullet(doc, "backend/src/websocket: autenticación del socket, eventos de chat, video y reportes.")
    add_bullet(doc, "backend/src/workers y datos/gravedad: clasificación concurrente basada en reglas CSV.")
    add_bullet(doc, "frontend/src/componentes, paginas, contextos y servicios: interfaz, estado y consumo distribuido.")
    add_bullet(doc, "data_science: dataset, notebook KPI, notebook predictivo y salidas JSON consumidas por el backend.")

    add_heading(doc, "6.4 Arquitectura del sistema", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(ARCH_PATH), width=Inches(6.45))
    add_figure_caption(doc, 1, "Arquitectura distribuida de la solución", "Elaboración propia a partir de los módulos implementados en el repositorio.")
    add_body(doc, "La API REST concentra operaciones transaccionales de autenticación y reportes. Socket.IO comparte el servidor HTTP y distribuye eventos a salas específicas. El Worker Thread desacopla la clasificación. Los notebooks producen resúmenes JSON; el backend los expone mediante endpoints protegidos y el frontend los representa en dashboards.")

    add_heading(doc, "6.5 Flujo principal de un reporte", 2)
    flow_num = create_decimal_numbering(doc)
    for step in [
        "El usuario autenticado envía tipo, descripción, ubicación y evidencia mediante POST /api/reportes.",
        "El servidor valida los datos y crea un reporte con estado pendiente en el repositorio en memoria.",
        "El servicio de asignación selecciona una unidad compatible y próxima.",
        "El servicio lanza un Worker Thread para clasificar la gravedad sin bloquear la respuesta HTTP.",
        "Socket.IO emite reporte:nuevo y, al concluir el worker, reporte:gravedadActualizada.",
        "El operador acepta el reporte, cambia su estado a en_camino y finalmente a resuelto.",
        "El mapa obtiene la geometría de OSRM; si falla, utiliza una línea directa de respaldo.",
        "Mientras el estado es en_camino, requestAnimationFrame interpola la posición del marcador hasta el destino.",
    ]:
        add_number(doc, step, flow_num)

    add_heading(doc, "6.6 Rutas y algoritmo de desplazamiento", 2)
    add_body(doc, "El sistema utiliza tres mecanismos complementarios. Haversine calcula distancias geográficas; Dijkstra determina el camino mínimo dentro del grafo vial local; y OSRM solicita una geometría ajustada a las calles. La animación no calcula la ruta: recorre las coordenadas ya obtenidas, interpolando segmentos con requestAnimationFrame y actualizando el marcador mediante setLatLng.")
    add_table(doc, ["Mecanismo", "Finalidad", "Resultado"], [
        ("Haversine", "Comparar distancias entre coordenadas.", "Distancia aproximada sobre la esfera terrestre."),
        ("Dijkstra", "Seleccionar la ruta mínima dentro del grafo definido.", "Secuencia de nodos y costo acumulado."),
        ("OSRM", "Ajustar el recorrido a calles transitables.", "Geometría detallada en coordenadas."),
        ("Interpolación", "Mover visualmente la unidad sobre la geometría.", "Posición animada y visible hasta la emergencia."),
    ], [1700, 3700, 3960])

    add_heading(doc, "7. Aplicación de los contenidos de la asignatura", 1)
    add_heading(doc, "7.1 Comunicación cliente-servidor", 2)
    add_body(doc, "Los servicios del frontend encapsulan llamadas fetch hacia /api. Vite actúa como proxy de desarrollo hacia el backend en el puerto 3001. Las rutas Express separan autenticación, usuarios, reportes, analítica y predicciones. Esta organización mantiene el contrato HTTP independiente de la interfaz.")

    add_heading(doc, "7.2 Concurrencia mediante Worker Threads", 2)
    add_body(doc, "La creación del reporte responde sin esperar que el hilo principal ejecute el análisis de texto. El servicio crea un Worker, publica los datos con postMessage y procesa el resultado mediante el evento message. El worker carga cuatro CSV de reglas y eleva la gravedad cuando encuentra términos con un peso superior.")
    add_code(doc, """
    const worker = new Worker(RUTA_WORKER);
    worker.postMessage({ id: reporte.id, tipo: reporte.tipo,
      descripcion: reporte.descripcion });
    worker.on('message', (resultado) => {
      const actualizado = reportesAlmacenamiento.actualizarGravedad(
        resultado.reporteId, resultado.gravedad);
      notificacionServicio.notificarGravedadActualizada(actualizado);
      worker.terminate();
    });
    """, "Fragmento 1. Delegación de la clasificación al Worker Thread")

    add_heading(doc, "7.3 WebSocket, chat y videollamada", 2)
    add_body(doc, "Socket.IO autentica la conexión y registra handlers independientes. El chat conserva hasta cien mensajes por sala en memoria, limita multimedia a 5 MB y emite indicadores de escritura. La videollamada utiliza WebRTC en topología mesh: Socket.IO solo transporta ofertas, respuestas y candidatos ICE; el audio y video circulan entre pares.")
    add_code(doc, """
    socket.on('video:oferta', ({ para, oferta }) => {
      socket.to(para).emit('video:oferta', {
        de: socket.id, nombre: socket.usuario.nombre, oferta
      });
    });
    socket.on('video:ice', ({ para, candidato }) => {
      socket.to(para).emit('video:ice', { de: socket.id, candidato });
    });
    """, "Fragmento 2. Señalización WebRTC sobre Socket.IO")

    add_heading(doc, "7.4 JWT y autorización por roles", 2)
    add_body(doc, "Después del login, el backend firma un JWT que contiene id, correo, rol, nombre y unidad. El middleware de autenticación valida el token antes de permitir operaciones protegidas. En el frontend, RutaProtegida restringe las vistas del operador y administrador.")
    add_code(doc, """
    const token = generarToken({
      id: usuario.id, correo: usuario.correo, rol: usuario.rol,
      nombre: usuario.nombre, unidadId: usuario.unidadId || null
    });
    """, "Fragmento 3. Generación de sesión JWT")

    add_heading(doc, "7.5 Google OAuth 2.0", 2)
    add_body(doc, "El botón carga Google Identity Services y obtiene una credencial en modo popup. El backend verifica el ID token con google-auth-library y exige que la audiencia coincida con GOOGLE_CLIENT_ID. Si la cuenta no existe, se crea con rol usuario; las cuentas de operador o administrador no pueden ingresar mediante este flujo. Para evitar origin_mismatch, el origen exacto, incluido esquema y puerto, debe estar registrado en Google Cloud.")
    add_code(doc, """
    const ticket = await googleClient.verifyIdToken({
      idToken: credential,
      audience: googleClientId
    });
    const payload = ticket.getPayload();
    """, "Fragmento 4. Validación del ID token de Google")

    add_heading(doc, "7.6 Manejo centralizado de errores", 2)
    add_body(doc, "ValidacionError, NoAutorizadoError, ProhibidoError y NoEncontradoError especializan fallos de dominio. El middleware final distingue errores controlados de excepciones inesperadas; registra el incidente y evita devolver el stack al cliente.")
    add_code(doc, """
    if (err instanceof AppError) {
      logger.warn('ERROR_HANDLER', err.message, {
        codigoEstado: err.codigoEstado, ruta: req.originalUrl
      });
      return res.status(err.codigoEstado).json({
        exito: false, mensaje: err.message
      });
    }
    """, "Fragmento 5. Respuesta uniforme para errores de dominio")

    add_heading(doc, "7.7 Logs y observabilidad", 2)
    add_body(doc, "El logger Singleton produce líneas con fecha ISO, nivel, módulo, mensaje y datos JSON opcionales. El middleware HTTP registra solicitudes y los servicios agregan eventos de negocio como creación de reportes, cambios de estado, conexión de sockets y fallos del worker. Queda pendiente persistir y centralizar estas señales.")

    add_heading(doc, "7.8 Consistencia y transiciones de estado", 2)
    add_body(doc, "La lógica de negocio limita las transiciones para evitar saltos arbitrarios. El flujo principal es pendiente → aceptado → en_camino → resuelto. El alias en_proceso se conserva por compatibilidad con versiones anteriores, pero la interfaz actual conduce al flujo de tres acciones operativas: aceptar, salir hacia el destino y finalizar.")
    add_code(doc, """
    const transicionesPermitidas = {
      pendiente: ['aceptado'],
      aceptado: ['en_camino'],
      en_camino: ['resuelto'],
      en_proceso: ['en_camino', 'resuelto'],
      resuelto: []
    };
    """, "Fragmento 6. Máquina de estados validada en el servicio")

    add_heading(doc, "7.9 Principios SOLID y patrones", 2)
    add_body(doc, "La separación entre rutas, controladores, servicios y almacenamiento aplica responsabilidad única. El frontend divide componentes visuales, contextos y servicios HTTP. El logger utiliza Singleton y los repositorios encapsulan el acceso a datos. La implementación no constituye una arquitectura completamente desacoplada por interfaces, pero presenta límites claros para sustituir almacenamiento o notificaciones.")

    add_heading(doc, "7.10 Transacciones", 2)
    add_body(doc, "No existe una base de datos transaccional en la versión actual. Las operaciones sobre Map() son síncronas dentro del proceso, pero no ofrecen durabilidad, rollback ni coordinación entre nodos. Por tanto, el requisito se considera parcial. Una implementación posterior debería utilizar PostgreSQL y una transacción para crear el reporte, registrar el estado inicial y asociar la unidad, complementada con un patrón outbox para publicar eventos sin perder consistencia.")

    add_heading(doc, "8. Módulo de ciencia de datos", 1)
    add_body(doc, "El notebook analisis_emergencias_kpi.ipynb procesa 24 154 registros y 43 variables del archivo mdi_detenidosaprehendidos_pm_2026_enero_abril.xlsx. Se obtuvieron 24 059 registros georreferenciados, equivalentes al 99,61 %. El dashboard del operador combina este resumen histórico con los reportes vivos almacenados por la aplicación.")
    add_table(doc, ["Indicador", "Resultado", "Interpretación"], [
        ("Registros", "24 154", "Volumen analizado entre enero y abril de 2026."),
        ("Georreferenciación", "99,61 %", "Cobertura suficiente para mapas, con revisión de faltantes."),
        ("Edad promedio", "32,4 años", "Descripción de la población registrada."),
        ("Registros con arma", "9,64 %", "Variable de severidad para segmentación."),
        ("Hora de mayor incidencia", "18:00", "Ventana histórica de atención prioritaria."),
        ("Provincia principal", "Guayas", "Mayor cantidad de registros del conjunto."),
    ], [2400, 1600, 5360])
    add_body(doc, "Debe evitarse interpretar este dataset como representación exacta de todas las emergencias. La fuente contiene detenidos y aprehendidos; por ello, los resultados describen registros policiales observados y funcionan como demostración analítica para el proyecto.")

    add_heading(doc, "9. Módulo de predicciones", 1)
    add_body(doc, "El notebook predicciones_emergencias.ipynb implementa un baseline estadístico interpretable. No utiliza una red neuronal ni promete causalidad. La zona se puntúa con volumen histórico, presencia de armas y ocurrencia nocturna; la probabilidad horaria y diaria se calcula a partir de frecuencias; y el forecast diario combina el promedio del día de la semana con una tendencia reciente.")
    add_table(doc, ["Salida", "Resultado principal", "Uso previsto"], [
        ("Zona de mayor riesgo", "Pascuales, Guayaquil; score 970", "Priorización preventiva y asignación de recursos."),
        ("Horario probable", "Pico operativo cercano a las 18:00", "Refuerzo de turnos y monitoreo."),
        ("Día con mayor incidencia", "Sábado; probabilidad relativa 0,1553", "Planificación semanal."),
        ("Forecast 2 de mayo", "233 incidentes estimados", "Referencia cuantitativa del baseline."),
    ], [2300, 3000, 4060])
    add_body(doc, "La predicción debe considerarse apoyo y no decisión automatizada. La salida fue generada con información histórica limitada a cuatro meses y requiere validación temporal, métricas de error y actualización periódica antes de un uso real.")

    add_heading(doc, "10. Estado actual del proyecto", 1)
    add_table(doc, ["Componente", "Estado", "Evidencia en el repositorio", "Observación"], [
        ("Interfaz y navegación", "Terminado", "App.jsx y páginas por rol", "Frontend compilado con Vite."),
        ("API REST", "Terminado", "Rutas, controladores y servicios", "Incluye endpoint de salud."),
        ("JWT y roles", "Terminado", "jwt.utilidad y middlewares", "Protección backend y frontend."),
        ("Google OAuth", "Parcial", "BotonGoogle y autenticacion.servicio", "Requiere origen autorizado y configuración externa."),
        ("Socket.IO y chat", "Terminado", "chat.handler y SocketContexto", "Historial temporal por sala."),
        ("Videollamada", "Terminado para demo", "VideollamadaChat y señalización", "Calidad depende de red; falta TURN propio."),
        ("Estados del reporte", "Terminado", "Servicio y botones de detalle", "Flujo pendiente/aceptado/en_camino/resuelto."),
        ("Ruta y animación", "Terminado", "MapaRuta.jsx", "Dijkstra, OSRM y fallback."),
        ("KPI", "Terminado", "Notebook, JSON y dashboard", "Combina datos históricos y vivos."),
        ("Predicciones", "Baseline terminado", "Notebook, JSON y dashboard", "Falta evaluación fuera de muestra."),
        ("Persistencia", "Pendiente", "Map() en memoria", "Los datos se pierden al reiniciar."),
        ("Blockchain", "Pendiente", "Sin módulos implementados", "Se evaluará como capa de trazabilidad."),
        ("Cola y clúster", "Pendiente", "Sin broker ni despliegue multi-nodo", "Trabajo de tercera unidad."),
    ], [1700, 1300, 3000, 3360])

    add_heading(doc, "10.1 Errores y riesgos identificados", 2)
    add_bullet(doc, "El origen OAuth debe coincidir exactamente con la URL del navegador; un cambio entre HTTP, HTTPS, localhost o ngrok produce origin_mismatch.")
    add_bullet(doc, "Los certificados autofirmados habilitan TLS, pero no generan confianza automática en otros dispositivos.")
    add_bullet(doc, "La videollamada mesh incrementa el consumo de ancho de banda conforme crece el número de participantes.")
    add_bullet(doc, "El almacenamiento en memoria impide recuperar reportes, mensajes y usuarios tras reiniciar el backend.")
    add_bullet(doc, "La API pública de OSRM y los túneles gratuitos pueden presentar latencia, límites o interrupciones externas.")
    add_bullet(doc, "No existe una suite automatizada que cubra transiciones, autorización, sockets y predicciones.")

    add_heading(doc, "11. Resultados actuales", 1)
    add_body(doc, "La revisión del código y la compilación del frontend permiten confirmar que la aplicación integra módulos distribuidos dentro de un mismo flujo. Los resultados cuantitativos de ciencia de datos provienen de los JSON generados por los notebooks. Las siguientes evidencias visuales deben reemplazarse por capturas obtenidas durante la ejecución final, con datos y fechas visibles.")
    add_evidence_placeholder(doc, 2, "Inicio de sesión y autenticación de usuario", "Insertar captura del login local o de Google después de autorizar el origen.")
    add_evidence_placeholder(doc, 3, "Centro de operaciones con reportes y mapa", "Insertar captura legible del panel del operador.")
    add_evidence_placeholder(doc, 4, "Cambio de estado y unidad en camino", "Insertar captura del estado en_camino con el marcador visible sobre la ruta.")
    add_evidence_placeholder(doc, 5, "Chat y videollamada por sala de emergencia", "Insertar captura con dos participantes conectados y controles de audio/video.")
    add_evidence_placeholder(doc, 6, "Dashboard de KPI y mapa de calor", "Insertar captura del módulo /operador/analitica.")
    add_evidence_placeholder(doc, 7, "Dashboard de predicciones", "Insertar captura del módulo /operador/predicciones.")
    add_evidence_placeholder(doc, 8, "Logs y excepción controlada", "Insertar terminal con timestamp, nivel, módulo y una respuesta de error validada.")

    add_heading(doc, "12. Resultados esperados", 1)
    add_bullet(doc, "Persistir usuarios, reportes, cambios de estado y auditoría en PostgreSQL mediante transacciones.")
    add_bullet(doc, "Publicar eventos de reportes en un broker para desacoplar notificaciones, analítica y trazabilidad.")
    add_bullet(doc, "Desplegar varias instancias del backend detrás de un balanceador con un adaptador compartido para Socket.IO.")
    add_bullet(doc, "Incorporar un servidor TURN para mejorar la conectividad WebRTC en redes restrictivas.")
    add_bullet(doc, "Registrar hashes de eventos críticos en blockchain sin exponer datos personales ni sustituir la base transaccional.")
    add_bullet(doc, "Evaluar el forecast con MAE o RMSE, validación temporal y reentrenamiento programado.")
    add_bullet(doc, "Centralizar logs, métricas y trazas con OpenTelemetry y alertas operativas.")

    add_heading(doc, "13. Pruebas propuestas", 1)
    add_table(doc, ["Prueba", "Procedimiento", "Resultado esperado"], [
        ("JWT válido", "Iniciar sesión y consultar una ruta protegida.", "200 y datos del usuario autorizado."),
        ("JWT ausente", "Consultar ruta protegida sin token.", "401 con mensaje controlado."),
        ("Rol insuficiente", "Usuario intenta cambiar estado.", "403 sin modificar el reporte."),
        ("Estado inválido", "Intentar pendiente → resuelto.", "400 y permanencia del estado original."),
        ("Concurrencia", "Crear múltiples reportes simultáneos.", "API receptiva y clasificación posterior por worker."),
        ("Socket", "Abrir usuario y operador en la misma emergencia.", "Mensajes y estados visibles en ambos clientes."),
        ("Ruta", "Cambiar a en_camino.", "Marcador recorre la ruta y llega al destino."),
        ("OAuth", "Ingresar desde un origen autorizado.", "Credencial validada y sesión JWT local."),
        ("Predicciones", "Consultar /api/predicciones/resumen.", "JSON protegido y dashboard renderizado."),
    ], [1800, 3800, 3760])

    add_heading(doc, "14. Conclusiones parciales", 1)
    conclusions_num = create_decimal_numbering(doc)
    add_number(doc, "La aplicación integra satisfactoriamente comunicación HTTP y Socket.IO, autenticación JWT, Google Identity Services, mapas, chat y procesamiento concurrente dentro de una arquitectura modular de frontend y backend.", conclusions_num)
    add_number(doc, "El Worker Thread permite separar la clasificación de gravedad del hilo principal, mientras que los eventos en tiempo real mantienen visibles los cambios de reporte; no obstante, la consistencia durable aún depende de incorporar una base de datos transaccional.", conclusions_num)
    add_number(doc, "Los módulos KPI y predictivo convierten un conjunto de 24 154 registros en información operativa, pero el forecast actual es un baseline interpretable que requiere validación temporal y no debe presentarse como inteligencia artificial de alta precisión.", conclusions_num)
    add_number(doc, "Los principales pendientes son persistencia, pruebas automatizadas, observabilidad centralizada, despliegue multi-nodo y trazabilidad blockchain. Estas mejoras permitirán pasar de una demostración académica funcional a una solución distribuida más resistente.", conclusions_num)

    add_heading(doc, "15. Referencias bibliográficas", 1)
    references = [
        "Google. (2025). Sign in with Google for Web. Google for Developers. https://developers.google.com/identity/gsi/web",
        "Internet Engineering Task Force. (2022). HTTP semantics (RFC 9110). RFC Editor. https://www.rfc-editor.org/rfc/rfc9110",
        "Internet Engineering Task Force. (2025). Best current practice for OAuth 2.0 security (RFC 9700). RFC Editor. https://www.rfc-editor.org/rfc/rfc9700",
        "Node.js Contributors. (2026). Worker threads. Node.js documentation. https://nodejs.org/api/worker_threads.html",
        "OpenTelemetry Authors. (2025). OpenTelemetry specification. https://opentelemetry.io/docs/specs/otel/",
        "React Team. (2025). React documentation: Describing the UI. https://react.dev/learn/describing-the-ui",
        "Socket.IO. (2024). Socket.IO documentation: Rooms and emitting events. https://socket.io/docs/v4/",
        "World Wide Web Consortium. (2025). WebRTC: Real-time communication in browsers. https://www.w3.org/TR/webrtc/",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(ref)
        set_run_font(r, size=10.5)

    doc.add_page_break()
    add_heading(doc, "16. Anexos", 1)
    add_heading(doc, "Anexo A. Instrucciones de ejecución", 2)
    add_code(doc, """
    # Terminal 1
    cd backend
    npm install
    npm run dev

    # Terminal 2
    cd frontend
    npm install
    npm run dev
    """)
    add_body(doc, "El backend utiliza el puerto 3001 y el frontend el puerto 5173. Si se usa HTTPS local, el navegador debe confiar en el certificado; para acceso externo se requiere un túnel activo y un origen OAuth autorizado.")

    add_heading(doc, "Anexo B. Variables de entorno", 2)
    add_code(doc, """
    # backend/.env.example
    PORT=3001
    JWT_SECRET=reemplace_con_un_secreto_seguro
    JWT_EXPIRATION=24h
    CORS_ORIGIN=https://localhost:5173
    GOOGLE_CLIENT_ID=su_client_id.apps.googleusercontent.com

    # frontend/.env.example
    VITE_GOOGLE_CLIENT_ID=su_client_id.apps.googleusercontent.com
    VITE_SOCKET_URL=
    """)
    add_body(doc, "El Client ID del frontend y el backend debe coincidir. Los secretos reales, claves privadas y certificados no deben publicarse en el repositorio.")

    add_heading(doc, "Anexo C. Lista de evidencias antes de exportar a PDF", 2)
    for item in [
        "Completar docente, NRC y periodo académico de la portada.",
        "Actualizar el índice general y el índice de figuras con Ctrl+A y F9 en Word.",
        "Reemplazar los siete recuadros de evidencia por capturas reales y legibles.",
        "Agregar enlace al repositorio y, si se solicita, enlace al video demostrativo.",
        "Verificar que el repositorio incluya README.md y .env.example sin secretos.",
        "Exportar el documento final como un único PDF y comprobar todas las páginas.",
    ]:
        add_bullet(doc, item)

    doc.core_properties.title = "Sistema distribuido de reporte, gestión y análisis de emergencias"
    doc.core_properties.subject = "Informe técnico - Aplicaciones Distribuidas"
    doc.core_properties.author = "Adonis Alegría"
    doc.core_properties.keywords = "emergencias, aplicaciones distribuidas, WebSocket, JWT, OAuth 2.0"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
